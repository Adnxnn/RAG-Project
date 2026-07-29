import hashlib, json
from pathlib import Path
import pandas as pd
import fitz
from docx import Document
from unstructured.partition.auto import partition
from app.core.models import KnowledgeBlock
from app.ingestion.vision import describe_image


def _id(seed: str) -> str:
    return hashlib.sha256(seed.encode()).hexdigest()[:24]


def load_file(path: Path) -> list[KnowledgeBlock]:
    suffix = path.suffix.lower()
    source_id = _id(str(path.resolve()))
    if suffix == ".pdf": return _pdf(path, source_id)
    if suffix in {".xlsx", ".xls"}: return _excel(path, source_id)
    if suffix == ".csv": return _csv(path, source_id)
    if suffix == ".docx": return _docx(path, source_id)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}: return _image(path, source_id)
    if suffix in {".txt", ".md", ".py", ".js", ".ts", ".tsx", ".java", ".sql", ".json", ".yaml", ".yml"}:
        return [KnowledgeBlock(id=_id(source_id+"0"), source_id=source_id, kind="code" if suffix not in {".txt", ".md"} else "text", text=path.read_text(errors="ignore"), metadata={"file_name": path.name, "locator": "file"})]
    return _unstructured(path, source_id)


def _pdf(path: Path, source_id: str) -> list[KnowledgeBlock]:
    blocks=[]
    doc=fitz.open(path)
    for i,page in enumerate(doc):
        text=page.get_text("text").strip()
        if text:
            blocks.append(KnowledgeBlock(id=_id(f"{source_id}:p:{i}"),source_id=source_id,kind="text",text=text,metadata={"file_name":path.name,"page":i+1,"locator":f"page {i+1}"}))
        for j,img in enumerate(page.get_images(full=True)):
            pix=fitz.Pixmap(doc,img[0]); data=pix.tobytes("png")
            description=describe_image(data, f"Image from {path.name}, page {i+1}")
            blocks.append(KnowledgeBlock(id=_id(f"{source_id}:p:{i}:img:{j}"),source_id=source_id,kind="image",text=description,raw={"page":i+1,"image_index":j},metadata={"file_name":path.name,"page":i+1,"locator":f"page {i+1}, image {j+1}"}))
    if not blocks:
        for e in partition(filename=str(path), strategy="hi_res"):
            blocks.append(KnowledgeBlock(id=_id(source_id+str(len(blocks))),source_id=source_id,kind="text",text=str(e),metadata={"file_name":path.name,"locator":"OCR/layout extraction"}))
    return blocks


def _excel(path: Path, source_id: str) -> list[KnowledgeBlock]:
    out=[]
    book=pd.ExcelFile(path)
    for sheet in book.sheet_names:
        df=pd.read_excel(path,sheet_name=sheet)
        text=f"Sheet: {sheet}\nColumns: {', '.join(map(str,df.columns))}\n{df.to_markdown(index=False)}"
        out.append(KnowledgeBlock(id=_id(source_id+sheet),source_id=source_id,kind="sheet",text=text,raw=df.where(pd.notna(df),None).to_dict("records"),metadata={"file_name":path.name,"sheet":sheet,"locator":f"sheet {sheet}, used range A1:{_excel_col(len(df.columns))}{len(df)+1}"}))
    return out


def _excel_col(n:int)->str:
    s=""
    while n: n,r=divmod(n-1,26); s=chr(65+r)+s
    return s or "A"


def _csv(path: Path, source_id: str) -> list[KnowledgeBlock]:
    df=pd.read_csv(path)
    return [KnowledgeBlock(id=_id(source_id),source_id=source_id,kind="table",text=df.to_markdown(index=False),raw=df.where(pd.notna(df),None).to_dict("records"),metadata={"file_name":path.name,"locator":f"rows 1-{len(df)+1}"})]


def _docx(path: Path, source_id: str) -> list[KnowledgeBlock]:
    d=Document(path); out=[]
    for i,p in enumerate(d.paragraphs):
        if p.text.strip(): out.append(KnowledgeBlock(id=_id(f"{source_id}:p:{i}"),source_id=source_id,kind="text",text=p.text,metadata={"file_name":path.name,"locator":f"paragraph {i+1}"}))
    for i,t in enumerate(d.tables):
        rows=[[c.text for c in r.cells] for r in t.rows]
        out.append(KnowledgeBlock(id=_id(f"{source_id}:t:{i}"),source_id=source_id,kind="table",text=json.dumps(rows,ensure_ascii=False),raw=rows,metadata={"file_name":path.name,"locator":f"table {i+1}"}))
    return out


def _image(path: Path, source_id: str) -> list[KnowledgeBlock]:
    return [KnowledgeBlock(id=_id(source_id),source_id=source_id,kind="image",text=describe_image(path.read_bytes(),path.name),metadata={"file_name":path.name,"locator":"image"})]


def _unstructured(path: Path, source_id: str) -> list[KnowledgeBlock]:
    return [KnowledgeBlock(id=_id(source_id+str(i)),source_id=source_id,kind="text",text=str(e),metadata={"file_name":path.name,"locator":f"element {i+1}"}) for i,e in enumerate(partition(filename=str(path),strategy="hi_res")) if str(e).strip()]
